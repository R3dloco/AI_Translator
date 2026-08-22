import os
import io
import json
import zipfile
from typing import List, Dict, Tuple, Union
import docx
from extractors.document_extractor import DocumentExtractor
from extractors.media_extractor import MediaExtractor
from extractors.ocr_extractor import OcrExtractor

class FileHandler:
    """Handles directory scanning, batch file operations, and saving translations to disk."""

    IGNORED_DIRS = {
        ".git", ".svn", ".hg", "venv", ".venv", "node_modules",
        "__pycache__", ".idea", ".vscode", "dist", "build", ".streamlit"
    }

    @staticmethod
    def scan_directory(dir_path: str, recursive: bool = True) -> List[str]:
        """Scans a local directory for supported document, image, audio, and video files."""
        if not dir_path or not dir_path.strip():
            return []
        
        expanded_path = os.path.expanduser(dir_path.strip())
        if not os.path.exists(expanded_path) or not os.path.isdir(expanded_path):
            return []

        supported_files = []
        doc_exts = set(DocumentExtractor.get_supported_extensions())
        media_exts = set(MediaExtractor.get_supported_extensions())
        ocr_exts = set(OcrExtractor.get_supported_extensions())
        all_exts = doc_exts.union(media_exts).union(ocr_exts)

        if recursive:
            for root, dirs, files in os.walk(expanded_path):
                # Filter out ignored directories in-place so os.walk does not recurse into them
                dirs[:] = [d for d in dirs if d not in FileHandler.IGNORED_DIRS and not d.startswith(".")]
                for file in files:
                    ext = os.path.splitext(file)[1].lower()
                    if ext in all_exts and not file.startswith("."):
                        supported_files.append(os.path.join(root, file))
        else:
            for file in os.listdir(expanded_path):
                full_path = os.path.join(expanded_path, file)
                if os.path.isfile(full_path):
                    ext = os.path.splitext(file)[1].lower()
                    if ext in all_exts and not file.startswith("."):
                        supported_files.append(full_path)

        return sorted(supported_files)

    @staticmethod
    def save_translation_to_file(
        original_filename: str,
        extracted_text: str,
        translated_text: str,
        output_dir: str,
        format_type: str = "txt",
        target_language: str = "English"
    ) -> str:
        """
        Saves translated result to disk in specified output_dir.
        `format_type` can be 'txt', 'md', 'docx', or 'json'.
        Returns the saved file path.
        """
        expanded_output_dir = os.path.expanduser(output_dir.strip()) if output_dir else "./translated_output"
        os.makedirs(expanded_output_dir, exist_ok=True)
        base_name = os.path.splitext(os.path.basename(original_filename))[0]
        lang_suffix = target_language.replace(" ", "_")
        output_path = os.path.join(expanded_output_dir, f"{base_name}_translated_{lang_suffix}.{format_type}")

        if format_type == "txt":
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(translated_text)
        elif format_type == "md":
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(f"# Translation ({target_language}): {base_name}\n\n")
                f.write(translated_text)
        elif format_type == "docx":
            doc = docx.Document()
            doc.add_heading(f"Translation ({target_language}): {base_name}", level=1)
            for paragraph in translated_text.split("\n\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            doc.save(output_path)
        elif format_type == "json":
            data = {
                "original_filename": os.path.basename(original_filename),
                "target_language": target_language,
                "extracted_text": extracted_text,
                "translated_text": translated_text
            }
            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
        elif format_type in ("srt", "vtt"):
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(translated_text)
        else:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(translated_text)

        return output_path

    @staticmethod
    def create_zip_archive(files_dict: Dict[str, Union[str, bytes]]) -> bytes:
        """
        Creates a ZIP archive in memory containing saved files.
        `files_dict` maps relative path -> file content string or bytes.
        """
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for filename, content in files_dict.items():
                data = content.encode("utf-8") if isinstance(content, str) else content
                zip_file.writestr(filename, data)
        zip_buffer.seek(0)
        return zip_buffer.getvalue()

